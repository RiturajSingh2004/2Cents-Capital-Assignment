import os
import re
from typing import Dict, List, Tuple, Optional
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn
import mammoth
import asyncio
from pathlib import Path

from app.models import DocumentType, DocumentFlag, FlagSeverity
from config import settings

class DocumentParser:
    def __init__(self):
        self.supported_formats = ['.docx']
    
    async def parse_document(self, file_path: str) -> Dict:
        """Parse uploaded document and extract content and metadata"""
        try:
            # Basic file validation
            if not self._validate_file(file_path):
                raise ValueError("Unsupported file format or file too large")
            
            # Extract content
            content = await self._extract_content(file_path)
            
            # Detect document type
            doc_type = self._detect_document_type(content['text'])
            
            # Extract structure
            structure = self._extract_structure(file_path)
            
            return {
                'file_path': file_path,
                'document_type': doc_type,
                'content': content,
                'structure': structure,
                'metadata': self._extract_metadata(file_path)
            }
            
        except Exception as e:
            raise Exception(f"Document parsing failed: {str(e)}")
    
    def _validate_file(self, file_path: str) -> bool:
        """Validate file format and size"""
        if not os.path.exists(file_path):
            return False
            
        file_ext = Path(file_path).suffix.lower()
        if file_ext not in self.supported_formats:
            return False
            
        file_size = os.path.getsize(file_path)
        if file_size > settings.MAX_FILE_SIZE:
            return False
            
        return True
    
    async def _extract_content(self, file_path: str) -> Dict:
        """Extract text content from document"""
        try:
            # Use python-docx for structure and mammoth for clean text
            doc = Document(file_path)
            
            # Extract paragraphs with formatting
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Get clean text using mammoth
            with open(file_path, 'rb') as doc_file:
                result = mammoth.extract_raw_text(doc_file)
                clean_text = result.value
            
            return {
                'text': clean_text,
                'paragraphs': paragraphs,
                'tables': tables,
                'word_count': len(clean_text.split())
            }
            
        except Exception as e:
            raise Exception(f"Content extraction failed: {str(e)}")
    
    def _detect_document_type(self, text: str) -> DocumentType:
        """Detect document type based on content analysis"""
        text_lower = text.lower()
        
        # Define patterns for each document type
        patterns = {
            DocumentType.MEMORANDUM: [
                'memorandum of association', 'company objects', 'share capital',
                'liability of members', 'registered office'
            ],
            DocumentType.ARTICLES: [
                'articles of association', 'board of directors', 'general meeting',
                'dividend', 'transfer of shares'
            ],
            DocumentType.APPLICATION: [
                'application for registration', 'company registration',
                'business license application', 'adgm registration'
            ],
            DocumentType.BOARD_RESOLUTION: [
                'board resolution', 'resolved that', 'board meeting',
                'directors present', 'resolution passed'
            ]
        }
        
        # Score each document type
        scores = {}
        for doc_type, keywords in patterns.items():
            score = sum(1 for keyword in keywords if keyword in text_lower)
            scores[doc_type] = score
        
        # Return highest scoring type
        if max(scores.values()) > 0:
            return max(scores, key=scores.get)
        
        return DocumentType.UNKNOWN
    
    def _extract_structure(self, file_path: str) -> Dict:
        """Extract document structure (headings, sections)"""
        try:
            doc = Document(file_path)
            structure = {
                'headings': [],
                'sections': [],
                'outline': []
            }
            
            current_section = None
            
            for para in doc.paragraphs:
                if para.style and 'Heading' in para.style.name:
                    level = int(para.style.name.split()[-1]) if para.style.name.split()[-1].isdigit() else 1
                    heading = {
                        'text': para.text,
                        'level': level,
                        'style': para.style.name
                    }
                    structure['headings'].append(heading)
                    
                    if level == 1:
                        if current_section:
                            structure['sections'].append(current_section)
                        current_section = {
                            'title': para.text,
                            'content': [],
                            'subsections': []
                        }
                elif current_section and para.text.strip():
                    current_section['content'].append(para.text)
            
            if current_section:
                structure['sections'].append(current_section)
            
            return structure
            
        except Exception as e:
            return {'headings': [], 'sections': [], 'outline': []}
    
    def _extract_metadata(self, file_path: str) -> Dict:
        """Extract document metadata"""
        try:
            doc = Document(file_path)
            props = doc.core_properties
            
            return {
                'title': props.title or '',
                'author': props.author or '',
                'subject': props.subject or '',
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None,
                'revision': props.revision or 1,
                'pages': self._count_pages(doc),
                'file_size': os.path.getsize(file_path)
            }
            
        except Exception as e:
            return {}
    
    def _count_pages(self, doc: Document) -> int:
        """Estimate page count"""
        # Simple estimation based on content
        total_text = sum(len(para.text) for para in doc.paragraphs)
        return max(1, total_text // 2500)  # Roughly 2500 chars per page
    
    async def add_comments_to_document(self, file_path: str, flags: List[DocumentFlag], output_path: str):
        """Add comments and highlights to document based on analysis flags"""
        try:
            doc = Document(file_path)
            
            # Process each flag and add comments
            for flag in flags:
                # Find relevant paragraphs
                target_paragraphs = self._find_paragraphs_by_content(doc, flag.location)
                
                for para in target_paragraphs:
                    # Highlight text based on severity
                    self._highlight_paragraph(para, flag.severity)
                    
                    # Add comment (simplified - in practice would use proper comment API)
                    comment_text = f"{flag.title}: {flag.description}"
                    if flag.suggested_fix:
                        comment_text += f"\nSuggested Fix: {flag.suggested_fix}"
                    
                    # Add comment as footnote or text box (simplified implementation)
                    self._add_comment_marker(para, comment_text)
            
            # Save modified document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error adding comments: {str(e)}")
            return False
    
    def _find_paragraphs_by_content(self, doc: Document, search_text: str) -> List:
        """Find paragraphs containing specific text"""
        matching_paras = []
        for para in doc.paragraphs:
            if search_text.lower() in para.text.lower():
                matching_paras.append(para)
        return matching_paras
    
    def _highlight_paragraph(self, paragraph, severity: FlagSeverity):
        """Highlight paragraph based on flag severity"""
        # Color coding: Critical=Red, Warning=Yellow, Info=Blue
        color_map = {
            FlagSeverity.CRITICAL: RGBColor(255, 200, 200),  # Light red
            FlagSeverity.WARNING: RGBColor(255, 255, 200),   # Light yellow
            FlagSeverity.INFO: RGBColor(200, 200, 255)       # Light blue
        }
        
        color = color_map.get(severity, RGBColor(255, 255, 255))
        
        # Apply highlighting to runs
        for run in paragraph.runs:
            try:
                # Set highlight color
                run.font.highlight_color = color
            except:
                pass  # Some formatting might not be supported
    
    def _add_comment_marker(self, paragraph, comment_text: str):
        """Add a comment marker to paragraph"""
        # Simplified comment implementation
        # In production, would use proper Word comments API
        marker_text = f" [COMMENT: {comment_text[:100]}...]"
        paragraph.add_run(marker_text).font.color.rgb = RGBColor(255, 0, 0)